import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

templates_dir = r"c:\Projects\__CSPM\src\templates\invoices"
os.makedirs(templates_dir, exist_ok=True)

# 1. Standard Elite DOCX (Regular)
doc = Document()
# Header
h = doc.add_heading('CORY SCHNEIDER LAW OFFICE', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_addr = doc.add_paragraph('14 Parsons Court | Thornhill, Ontario | L4J 6Z4\n(416) 725-9364\tcoryschneiderlaw.ca')

# Summary Box (Using simple paragraphs with tabs for simplicity in DOCX)
doc.add_paragraph('Client\t\t\t{{ client_name }}')
doc.add_paragraph('Invoice #\t\t{{ invoice_number }}')
doc.add_paragraph('Date\t\t\t{{ date }}')
doc.add_paragraph('Fees\t\t\t${{ "%.2f"|format(total_fees) }}')
doc.add_paragraph('{% p if discount_amount > 0 %}Discount\t\t-${{ "%.2f"|format(discount_amount) }}{% p endif %}')
doc.add_paragraph('Total Tax\t\t${{ "%.2f"|format(total_tax) }}')
doc.add_paragraph('BALANCE DUE\t\t${{ "%.2f"|format(total_due) }}').bold = True

doc.add_paragraph('\nTO:\t{{ client_name }}\nc/o:\t{{ client_address }}\n')

doc.add_paragraph('{% p for matter in matters %}')
doc.add_paragraph('Work for: {{ matter.name }}').bold = True

table = doc.add_table(rows=2, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'DATE'
hdr_cells[1].text = 'DESCRIPTION'
hdr_cells[2].text = 'TIME (hrs)'
hdr_cells[3].text = 'FEE'

row_cells = table.rows[1].cells
row_cells[0].text = '{% tr for item in matter.line_items %}{{ item.date }}'
row_cells[1].text = '{{ item.description }}'
row_cells[2].text = '{{ "%.1f"|format(item.hours) }}'
row_cells[3].text = '${{ "%.2f"|format(item.amount) }}{% endtr %}'

doc.add_paragraph('\n{{ matter.name }} Sub-total Fees:\t\t${{ "%.2f"|format(matter.total_fees) }}')
doc.add_paragraph('{{ matter.name }} Sub-total HST:\t\t${{ "%.2f"|format(matter.total_tax) }}')
doc.add_paragraph('{% p endfor %}')

doc.add_paragraph('\nTOTAL FEES:\t\t${{ "%.2f"|format(total_fees - discount_amount) }}').bold = True
doc.add_paragraph('TOTAL TAX:\t\t${{ "%.2f"|format(total_tax) }}').bold = True
doc.add_paragraph('TOTAL DUE:\t\t${{ "%.2f"|format(total_due) }}').bold = True

doc.add_paragraph('\nHST Registration No. 78621 8222 RT0001\n')
doc.add_paragraph('Invoice Balance Due Upon Receipt\n').bold = True

doc.add_paragraph('This is my account herein\n')
doc.add_paragraph('Cory Schneider')
doc.add_paragraph('\nPLEASE MAKE PAYMENT PAYABLE TO "CORY SCHNEIDER".\nPreferred mode of payment: Interac e-transfer to cory@coryschneiderlaw.ca.\n')

doc.save(os.path.join(templates_dir, 'Standard_Elite.docx'))


# 2. LIHDC Format DOCX
doc2 = Document()
h2 = doc2.add_heading('CORY SCHNEIDER', level=1)
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_addr2 = doc2.add_paragraph('14 Parsons Court \u2666 Thornhill, Ontario \u2666 L4J 6Z4\n(416) 725-9364\tcory.schneider@lawyersinhouse.com')
p_addr2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc2.add_paragraph('Client\t\t\t{{ client_name }}')
doc2.add_paragraph('Invoice #\t\t{{ invoice_number }}')
doc2.add_paragraph('Date\t\t\t{{ date }}')
doc2.add_paragraph('Fees (to CS)\t\t${{ "%.2f"|format(total_fees) }}')
doc2.add_paragraph('Total Tax\t\t${{ "%.2f"|format(total_tax) }}')
doc2.add_paragraph('BALANCE DUE\t\t${{ "%.2f"|format(total_due) }}').bold = True

doc2.add_paragraph('\nTO: LIHDC PROFESSIONAL CORPORATION').bold = True
doc2.add_paragraph('Legal Fees').bold = True

table2 = doc2.add_table(rows=2, cols=5)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'DATE'
hdr_cells2[1].text = 'DESCRIPTION'
hdr_cells2[2].text = 'TIME (hrs)'
hdr_cells2[3].text = 'Amount (to LIH)'
hdr_cells2[4].text = 'Amount (to CS)'

row_cells2 = table2.rows[1].cells
row_cells2[0].text = '{% tr for item in line_items %}{{ item.date }}'
row_cells2[1].text = '{{ item.description }}'
row_cells2[2].text = '{{ "%.1f"|format(item.hours) }}'
row_cells2[3].text = '${{ "%.2f"|format(item.amount / 0.9) }}'
row_cells2[4].text = '${{ "%.2f"|format(item.amount) }}{% endtr %}'

doc2.add_paragraph('\nSub-total Fees (to CS):\t\t${{ "%.2f"|format(total_fees) }}')
doc2.add_paragraph('Sub-total HST:\t\t${{ "%.2f"|format(total_tax) }}')

doc2.add_paragraph('\nTotal Time (hrs):\t{{ "%.1f"|format(total_hours) }}')
doc2.add_paragraph('Rate ($/hr):\t${{ "%.2f"|format(rate) }}')
doc2.add_paragraph('Sub-Total Charge to Client:\t${{ "%.2f"|format(total_fees / 0.9) }}')
doc2.add_paragraph('Sub-Total Charge to LIH with 10.0% Deduction:\t${{ "%.2f"|format(total_fees) }}')
doc2.add_paragraph('Total Tax:\t${{ "%.2f"|format(total_tax) }}')
doc2.add_paragraph('Total:\t${{ "%.2f"|format(total_due) }}')

doc2.add_paragraph('\nHST Registration No. 78621 8222 RT0001\n')
doc2.add_paragraph('Invoice Balance Due Upon Receipt\n').bold = True

doc2.add_paragraph('This is my account herein\n')
doc2.add_paragraph('Cory Schneider')
doc2.add_paragraph('\nPLEASE MAKE PAYMENT PAYABLE TO "CORY SCHNEIDER". In accordance with s. 33 of the Solicitors Act (Ontario)....')

doc2.save(os.path.join(templates_dir, 'LIHDC_Format.docx'))

print("Updated DOCX templates generated.")
