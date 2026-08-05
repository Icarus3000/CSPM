from __future__ import annotations

from collections import Counter, defaultdict, deque
from datetime import date, datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
import csv
import json
import os
import re

from openpyxl import load_workbook

try:
    from docx import Document
    from docx.shared import Pt
except Exception:  # pragma: no cover - fallback for machines without python-docx
    Document = None
    Pt = None


BASE = Path(__file__).resolve().parents[1]


def configured_path(env_name: str, default: Path) -> Path:
    raw_value = os.environ.get(env_name, "").strip()
    if not raw_value:
        return default
    return Path(raw_value).expanduser().resolve()


SOURCE = configured_path("CSPM_RECON_SOURCE", BASE / "data" / "Dockets.xlsm")
TARGET = configured_path("CSPM_RECON_TARGET", BASE / "data" / "CSPM.xlsm")
SQLITE = BASE / "CSPM_Database.db"
OUT_DIR = configured_path("CSPM_RECON_OUT_DIR", BASE / "outputs" / "reports")
OUT_DIR.mkdir(parents=True, exist_ok=True)
STAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
DOCX_PATH = OUT_DIR / f"dockets_cspm_reconciliation_{STAMP}.docx"
MD_PATH = OUT_DIR / f"dockets_cspm_reconciliation_{STAMP}.md"
CSV_PATH = OUT_DIR / f"dockets_cspm_reconciliation_{STAMP}_discrepancies.csv"
JSON_PATH = OUT_DIR / f"dockets_cspm_reconciliation_{STAMP}_summary.json"
MISSING_SHEETS: list[dict[str, str]] = []


def clean(value):
    if value is None:
        return ""
    if isinstance(value, str):
        return re.sub(r"\s+", " ", value.strip())
    return value


def text(value) -> str:
    value = clean(value)
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, Decimal):
        return f"{value:,.2f}"
    return "" if value is None else str(value)


def norm_text(value) -> str:
    return re.sub(r"\s+", " ", str(clean(value))).casefold()


def norm_name(value) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(clean(value)).casefold())


def date_key(value) -> str:
    value = clean(value)
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    return str(value)[:10] if value else ""


def dec(value, places: str = "0.01") -> Decimal:
    quant = Decimal(places)
    if value is None or value == "":
        return Decimal("0").quantize(quant)
    if isinstance(value, str):
        value = value.replace("$", "").replace(",", "").strip()
        if not value:
            return Decimal("0").quantize(quant)
    try:
        return Decimal(str(value)).quantize(quant, rounding=ROUND_HALF_UP)
    except (InvalidOperation, ValueError):
        return Decimal("0").quantize(quant)


def money(value) -> str:
    return f"${dec(value):,.2f}"


def desc(value) -> str:
    return re.sub(r"\s*\[[^\]]+\]\s*$", "", str(clean(value)))


def is_blank(value) -> bool:
    return clean(value) in ("", None)


def is_actual_invoice(value) -> bool:
    return bool(re.match(r"^\d{2}-\d{4}(?:-[A-Z])?$", str(clean(value))))


def invoice_bucket(value) -> str:
    value = str(clean(value))
    upper = value.upper()
    if not value:
        return "Blank / unassigned"
    if is_actual_invoice(value):
        return "Invoice number"
    if upper == "BILLED":
        return "Legacy BILLED marker"
    if "HOLD" in upper or "FORGOT" in upper:
        return "Hold / bill later marker"
    if "FREE" in upper or "DO NOT BILL" in upper or "WRITE OFF" in upper:
        return "No-bill / write-off marker"
    return "Other marker"


def read_rows(path: Path, sheet: str, key_headers: list[str] | None = None) -> list[dict]:
    wb = load_workbook(path, data_only=True, keep_vba=True, read_only=False)
    try:
        try:
            ws = wb[sheet]
        except KeyError:
            MISSING_SHEETS.append({"path": str(path), "sheet": sheet})
            return []
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return []
        headers = [str(v).strip() if v is not None else "" for v in rows[0]]
        result = []
        for row_num, values in enumerate(rows[1:], start=2):
            row = {}
            for idx, header in enumerate(headers):
                if header:
                    row[header] = values[idx] if idx < len(values) else None
            meaningful = any(not is_blank(row.get(k)) for k in key_headers) if key_headers else any(
                not is_blank(v) for v in row.values()
            )
            if meaningful:
                row["_row"] = row_num
                result.append(row)
        return result
    finally:
        wb.close()


def add_record(records, category, severity, issue, source_row="", cspm_row="", identifier="", source="", cspm="", note=""):
    records.append(
        {
            "Category": category,
            "Severity": severity,
            "Issue": issue,
            "SourceRow": text(source_row),
            "CSPMRow": text(cspm_row),
            "Identifier": text(identifier),
            "SourceValue": text(source),
            "CSPMValue": text(cspm),
            "Note": text(note),
        }
    )


def docket_source_sig(row, include_invoice=False):
    sig = (
        date_key(row.get("Date")),
        desc(row.get("Description")),
        dec(row.get("Time (in hrs)")),
        dec(row.get("Hourly Rate/Flat Fee")),
        dec(row.get("Percentage")),
        dec(row.get("Amount to CS")),
    )
    return sig + ((clean(row.get("Invoice")),) if include_invoice else ())


def docket_target_sig(row, include_invoice=False):
    sig = (
        date_key(row.get("Date")),
        desc(row.get("Description")),
        dec(row.get("Time (in hrs) or Units")),
        dec(row.get("Hourly Rate/Flat Rate")),
        dec(row.get("Percentage")),
        dec(row.get("Amount to CS")),
    )
    return sig + ((clean(row.get("Invoice #")),) if include_invoice else ())


def timeentry_sig(row):
    return (
        date_key(row.get("Date")),
        desc(row.get("Description")),
        dec(row.get("Hours")),
        dec(row.get("ClientRate")),
        dec(row.get("SharePct")),
        dec(row.get("AmountToYou")),
    )


def summarize_dockets(rows, schema):
    if schema == "time":
        return {
            "Rows": len(rows),
            "Hours": sum(dec(r.get("Hours")) for r in rows),
            "AmountToCS": sum(dec(r.get("AmountToYou")) for r in rows),
            "TotalInclHST": sum(dec(r.get("TotalInclHST")) for r in rows),
            "InvoiceRows": "n/a",
            "WipCandidateRows": sum(1 for r in rows if str(clean(r.get("Status"))).casefold() != "billed"),
            "WipCandidateAmount": sum(
                dec(r.get("AmountToYou")) for r in rows if str(clean(r.get("Status"))).casefold() != "billed"
            ),
        }
    if schema == "source":
        hours, rate, amount, total, inv = (
            "Time (in hrs)",
            "Hourly Rate/Flat Fee",
            "Amount to CS",
            "Total Inclusive of HST",
            "Invoice",
        )
    else:
        hours, rate, amount, total, inv = (
            "Time (in hrs) or Units",
            "Hourly Rate/Flat Rate",
            "Amount to CS",
            "Total Inclusive of HST",
            "Invoice #",
        )
    return {
        "Rows": len(rows),
        "Hours": sum(dec(r.get(hours)) for r in rows),
        "AmountToCS": sum(dec(r.get(amount)) for r in rows),
        "TotalInclHST": sum(dec(r.get(total)) for r in rows),
        "InvoiceRows": sum(1 for r in rows if is_actual_invoice(r.get(inv))),
        "LegacyBilledRows": sum(1 for r in rows if str(clean(r.get(inv))).upper() == "BILLED"),
        "WipCandidateRows": sum(
            1 for r in rows if invoice_bucket(r.get(inv)) in {"Blank / unassigned", "Hold / bill later marker"}
        ),
        "WipCandidateAmount": sum(
            dec(r.get(amount))
            for r in rows
            if invoice_bucket(r.get(inv)) in {"Blank / unassigned", "Hold / bill later marker"}
        ),
        "RateTotalCheck": sum(dec(r.get(hours)) * dec(r.get(rate)) for r in rows),
    }


def disb_source_sig(row):
    return (
        date_key(row.get("Date")),
        clean(row.get("Client")),
        clean(row.get("Sub-Client")),
        desc(row.get("Description")),
        dec(row.get("Amount")),
        clean(row.get("Tax Exempt? (Y/N)")),
        dec(row.get("Bill %")),
        clean(row.get("Invoice")),
    )


def disb_target_sig(row):
    return (
        date_key(row.get("Date")),
        clean(row.get("ClientName")),
        clean(row.get("SubClient")),
        desc(row.get("Description")),
        dec(row.get("Amount")),
        clean(row.get("TaxExempt")),
        dec(row.get("BillPct")),
        clean(row.get("InvoiceRef")),
    )


def summarize_disb(rows, schema):
    invoice = "Invoice" if schema == "source" else "InvoiceRef"
    return {
        "Rows": len(rows),
        "Amount": sum(dec(r.get("Amount")) for r in rows),
        "InvoiceRows": sum(1 for r in rows if is_actual_invoice(r.get(invoice))),
        "WipCandidateAmount": sum(dec(r.get("Amount")) for r in rows if not is_actual_invoice(r.get(invoice))),
    }


def summarize_ledger(rows, schema):
    if schema == "source":
        fields = {
            "Billings": "Billings (excl. HST)",
            "HSTCollected": "HST Collected",
            "Expenses": "Expenses (excl. HST)",
            "HSTPaid": "HST Paid",
            "Collected": "Collected",
            "WriteOff": "Write Off",
            "Receivable": "Receivable",
        }
    else:
        fields = {
            "Billings": "BillingsExclHST",
            "HSTCollected": "HSTCollected",
            "Expenses": "ExpensesExclHST",
            "HSTPaid": "HSTPaid",
            "Collected": "Collected",
            "WriteOff": "WriteOff",
            "Receivable": "Receivable",
        }
    return {"Rows": len(rows), **{name: sum(dec(r.get(field)) for r in rows) for name, field in fields.items()}}


def ledger_key(row, schema):
    trx = clean(row.get("TrxID"))
    if trx:
        return ("trx", trx)
    ref = clean(row.get("Reference"))
    client = row.get("Client/Vendor" if schema == "source" else "ClientVendor")
    if ref:
        return ("ref", ref, date_key(row.get("Date")), norm_text(row.get("Category")), norm_text(client))
    if schema == "source":
        return (
            "sig",
            date_key(row.get("Date")),
            norm_text(client),
            norm_text(row.get("Description")),
            norm_text(row.get("Category")),
            dec(row.get("Billings (excl. HST)")),
            dec(row.get("Expenses (excl. HST)")),
            dec(row.get("Collected")),
            dec(row.get("Receivable")),
        )
    return (
        "sig",
        date_key(row.get("Date")),
        norm_text(client),
        norm_text(row.get("Description")),
        norm_text(row.get("Category")),
        dec(row.get("BillingsExclHST")),
        dec(row.get("ExpensesExclHST")),
        dec(row.get("Collected")),
        dec(row.get("Receivable")),
    )


def group_receivables(rows, schema):
    grouped = defaultdict(list)
    for row in rows:
        inv = clean(row.get("InvoiceNum"))
        if inv:
            grouped[inv].append(row)
    out = {}
    for inv, items in grouped.items():
        if schema == "source":
            total, paid, credits, balance = "Total_Invoiced", "Amount_Paid", "Credits/Adj", "Balance_Due"
        else:
            total, paid, credits, balance = "TotalInvoiced", "AmountPaid", "CreditsAdj", "BalanceDue"
        out[inv] = {
            "rows": items,
            "client": "; ".join(sorted({text(r.get("Client")) for r in items if text(r.get("Client"))})),
            "total": sum(dec(r.get(total)) for r in items),
            "paid": sum(dec(r.get(paid)) for r in items),
            "credits": sum(dec(r.get(credits)) for r in items),
            "balance": sum(dec(r.get(balance)) for r in items),
            "statuses": "; ".join(sorted({text(r.get("Status")) for r in items if text(r.get("Status"))})),
        }
    return out


def group_invoice_log(rows, schema):
    grouped = defaultdict(list)
    for row in rows:
        inv = clean(row.get("Invoice #" if schema == "source" else "InvoiceNum"))
        if inv:
            grouped[inv].append(row)
    out = {}
    for inv, items in grouped.items():
        if schema == "source":
            client, sub, date_col = "Client Name", "Sub-Client", "Invoice Date"
            fees, disb, tax, aggregate = "Total Fees", "Total Disbursements", "Total Tax", "Aggregate Billed to Client"
        else:
            client, sub, date_col = "ClientName", "SubClient", "InvoiceDate"
            fees, disb, tax, aggregate = "TotalFees", "TotalDisbursements", "TotalTax", "AggregateBilled"
        out[inv] = {
            "rows": items,
            "client": "; ".join(sorted({text(r.get(client)) for r in items if text(r.get(client))})),
            "subclient": "; ".join(sorted({text(r.get(sub)) for r in items if text(r.get(sub))})),
            "dates": "; ".join(sorted({date_key(r.get(date_col)) for r in items if date_key(r.get(date_col))})),
            "fees": sum(dec(r.get(fees)) for r in items),
            "disb": sum(dec(r.get(disb)) for r in items),
            "tax": sum(dec(r.get(tax)) for r in items),
            "aggregate": sum(dec(r.get(aggregate)) for r in items),
        }
    return out


def compare() -> tuple[dict, list[dict]]:
    MISSING_SHEETS.clear()
    src = {
        "Dockets": read_rows(SOURCE, "Dockets", ["Date", "Client", "Sub-Client", "Description", "Time (in hrs)", "Amount to CS", "Invoice"]),
        "Disbursements": read_rows(SOURCE, "Disbursements", ["Date", "Client", "Sub-Client", "Description", "Amount", "Invoice"]),
        "Clients": read_rows(SOURCE, "Clients", ["Client Name", "Client_ID"]),
        "Matters": read_rows(SOURCE, "Matters", ["Matter_ID", "Client_ID", "Description (Internal)"]),
        "Ledger": read_rows(SOURCE, "Ledger", ["Date", "Client/Vendor", "Description", "Reference", "Billings (excl. HST)", "Expenses (excl. HST)", "Collected", "Receivable"]),
        "Receivables": read_rows(SOURCE, "Receivables", ["InvoiceNum", "Date", "Client", "Total_Invoiced", "Balance_Due"]),
        "Invoice Log": read_rows(SOURCE, "Invoice Log", ["Invoice #", "Client Name", "Invoice Date", "Aggregate Billed to Client"]),
        "HST_Log": read_rows(SOURCE, "HST_Log"),
    }
    tgt = {
        "Dockets": read_rows(TARGET, "Dockets", ["Date", "Client", "Parent", "Description", "Time (in hrs) or Units", "Amount to CS", "Invoice #"]),
        "TimeEntries": read_rows(TARGET, "TimeEntries", ["EntryID", "Date", "ClientID", "Description", "Hours", "AmountToYou", "Status"]),
        "Disbursements": read_rows(TARGET, "Disbursements", ["Date", "ClientName", "SubClient", "Description", "Amount", "InvoiceRef"]),
        "Clients": read_rows(TARGET, "Clients", ["ClientID", "ClientName"]),
        "ClientProfiles": read_rows(TARGET, "ClientProfiles", ["ClientID", "LegalName", "DisplayName"]),
        "Parents": read_rows(TARGET, "Parents", ["ParentID", "ParentName"]),
        "Matters": read_rows(TARGET, "Matters", ["MatterID", "MatterNumber", "ClientID", "MatterName"]),
        "Ledger": read_rows(TARGET, "Ledger", ["Date", "ClientVendor", "Description", "Reference", "BillingsExclHST", "ExpensesExclHST", "Collected", "Receivable"]),
        "Receivables": read_rows(TARGET, "Receivables", ["InvoiceNum", "Date", "Client", "TotalInvoiced", "BalanceDue"]),
        "InvoiceLog": read_rows(TARGET, "InvoiceLog", ["InvoiceNum", "ClientName", "InvoiceDate", "AggregateBilled"]),
        "Transactions": read_rows(TARGET, "Transactions"),
        "HSTLog": read_rows(TARGET, "HSTLog"),
    }
    records = []
    for missing in MISSING_SHEETS:
        workbook_label = "CSPM" if Path(missing["path"]) == TARGET else "Dockets"
        add_record(
            records,
            "Workbook structure",
            "High",
            f"{workbook_label} workbook is missing required reconciliation sheet '{missing['sheet']}'",
            "",
            "",
            missing["sheet"],
            "",
            missing["path"],
            "The reconciliation continues with this sheet treated as empty.",
        )

    malformed_source_ledger = [
        row
        for row in src["Ledger"]
        if not any(
            clean(row.get(field))
            for field in ["Date", "Client/Vendor", "Description", "Category", "Reference", "TrxID", "ExternalRefID"]
        )
        and any(not is_blank(value) for key, value in row.items() if key != "_row")
    ]
    for row in malformed_source_ledger:
        nonblank = {
            key: text(value)
            for key, value in row.items()
            if key != "_row" and not is_blank(value)
        }
        add_record(
            records,
            "Ledger/A/R",
            "High",
            "Malformed source ledger row has values but no date/client/description/category/reference",
            row.get("_row"),
            "",
            "Malformed source ledger row",
            nonblank,
            "",
            "Excluded from ledger total and row-by-row reconciliation because it cannot be keyed safely.",
        )
    if malformed_source_ledger:
        src["Ledger"] = [row for row in src["Ledger"] if row not in malformed_source_ledger]
    source_receivable_rollups = [
        row for row in src["Receivables"] if norm_text(row.get("InvoiceNum")) in {"total", "totals"}
    ]
    target_receivable_rollups = [
        row for row in tgt["Receivables"] if norm_text(row.get("InvoiceNum")) in {"total", "totals"}
    ]
    if source_receivable_rollups:
        src["Receivables"] = [row for row in src["Receivables"] if row not in source_receivable_rollups]
    if target_receivable_rollups:
        tgt["Receivables"] = [row for row in tgt["Receivables"] if row not in target_receivable_rollups]

    # Docketed time and fees.
    target_by_core = defaultdict(deque)
    for row in tgt["Dockets"]:
        target_by_core[docket_target_sig(row)].append(row)
    matched = []
    missing_dockets = []
    for row in src["Dockets"]:
        key = docket_source_sig(row)
        if target_by_core[key]:
            matched.append((row, target_by_core[key].popleft()))
        else:
            missing_dockets.append(row)
    extra_dockets = [row for queue in target_by_core.values() for row in queue]
    for row in missing_dockets:
        add_record(
            records,
            "Docketed time/fees",
            "High",
            "Source docket row not found in CSPM by date, description, hours, rate, percent, and Amount to CS",
            row.get("_row"),
            "",
            f"{date_key(row.get('Date'))} | {clean(row.get('Client'))} | {desc(row.get('Description'))[:90]}",
            f"Hours {dec(row.get('Time (in hrs)'))}; amount {money(row.get('Amount to CS'))}; marker {text(row.get('Invoice'))}",
            "",
            "Likely missing or altered source time entry.",
        )
    for row in extra_dockets:
        add_record(
            records,
            "Docketed time/fees",
            "Medium",
            "CSPM docket row not present in Dockets.xlsm by core commercial fields",
            "",
            row.get("_row"),
            f"{date_key(row.get('Date'))} | {clean(row.get('Client'))} | {desc(row.get('Description'))[:90]}",
            "",
            f"Hours {dec(row.get('Time (in hrs) or Units'))}; amount {money(row.get('Amount to CS'))}; marker {text(row.get('Invoice #'))}",
            "May be legitimate CSPM-only work after the source workbook was last updated, or import divergence.",
        )

    invoice_changes = []
    client_label_changes = []
    for srow, trow in matched:
        if clean(srow.get("Invoice")) != clean(trow.get("Invoice #")):
            invoice_changes.append((srow, trow))
            add_record(
                records,
                "Docketed time/fees",
                "Medium",
                "Matching docket row has different invoice/status marker",
                srow.get("_row"),
                trow.get("_row"),
                f"{date_key(srow.get('Date'))} | {desc(srow.get('Description'))[:90]}",
                f"marker {text(srow.get('Invoice'))}",
                f"marker {text(trow.get('Invoice #'))}",
                "The time exists, but billed/WIP classification differs.",
            )
        if (clean(srow.get("Client")), clean(srow.get("Sub-Client"))) != (clean(trow.get("Parent")), clean(trow.get("Client"))):
            client_label_changes.append((srow, trow))

    time_source_counter = Counter(docket_source_sig(r) for r in src["Dockets"])
    time_target_counter = Counter(timeentry_sig(r) for r in tgt["TimeEntries"])

    # Disbursements.
    disb_missing = list((Counter(disb_source_sig(r) for r in src["Disbursements"]) - Counter(disb_target_sig(r) for r in tgt["Disbursements"])).elements())
    disb_extra = list((Counter(disb_target_sig(r) for r in tgt["Disbursements"]) - Counter(disb_source_sig(r) for r in src["Disbursements"])).elements())
    for key in disb_missing:
        add_record(records, "Disbursements", "High", "Source disbursement not found in CSPM", "", "", " | ".join(map(text, key)), text(key), "", "")
    for key in disb_extra:
        add_record(records, "Disbursements", "Medium", "CSPM disbursement not found in source", "", "", " | ".join(map(text, key)), "", text(key), "")

    # Ledger.
    src_ledger, tgt_ledger = defaultdict(deque), defaultdict(deque)
    for row in src["Ledger"]:
        src_ledger[ledger_key(row, "source")].append(row)
    for row in tgt["Ledger"]:
        tgt_ledger[ledger_key(row, "target")].append(row)
    ledger_mismatch, ledger_source_only, ledger_target_only = [], [], []
    num_pairs = [
        ("Billings (excl. HST)", "BillingsExclHST"),
        ("HST Collected", "HSTCollected"),
        ("Expenses (excl. HST)", "ExpensesExclHST"),
        ("HST Paid", "HSTPaid"),
        ("Collected", "Collected"),
        ("Write Off", "WriteOff"),
        ("Receivable", "Receivable"),
        ("OriginalAmount", "OriginalAmount"),
    ]
    text_pairs = [
        ("Date", "Date"),
        ("Client/Vendor", "ClientVendor"),
        ("Description", "Description"),
        ("Category", "Category"),
        ("Reference", "Reference"),
        ("ExternalRefID", "ExternalRefID"),
    ]
    for key, queue in src_ledger.items():
        peer = tgt_ledger.get(key, deque())
        while queue and peer:
            srow, trow = queue.popleft(), peer.popleft()
            diffs = []
            for sfield, tfield in num_pairs:
                if dec(srow.get(sfield)) != dec(trow.get(tfield)):
                    diffs.append(f"{sfield}: source {money(srow.get(sfield))} vs CSPM {money(trow.get(tfield))}")
            for sfield, tfield in text_pairs:
                sval = date_key(srow.get(sfield)) if sfield == "Date" else clean(srow.get(sfield))
                tval = date_key(trow.get(tfield)) if tfield == "Date" else clean(trow.get(tfield))
                if norm_text(sval) != norm_text(tval):
                    diffs.append(f"{sfield}: source {text(sval)} vs CSPM {text(tval)}")
            if diffs:
                ledger_mismatch.append((key, srow, trow, diffs))
        while queue:
            ledger_source_only.append((key, queue.popleft()))
    for key, queue in tgt_ledger.items():
        while queue:
            ledger_target_only.append((key, queue.popleft()))
    for key, srow, trow, diffs in ledger_mismatch:
        add_record(records, "Ledger/A/R", "High", "Ledger key exists in both files but values differ", srow.get("_row"), trow.get("_row"), key, "; ".join(diffs), "", "")
    for key, row in ledger_source_only:
        add_record(records, "Ledger/A/R", "High", "Source ledger row not found in CSPM", row.get("_row"), "", key, f"{date_key(row.get('Date'))} {text(row.get('Client/Vendor'))} {money(row.get('Receivable'))}", "", "")
    for key, row in ledger_target_only:
        add_record(records, "Ledger/A/R", "Medium", "CSPM ledger row not found in source", "", row.get("_row"), key, "", f"{date_key(row.get('Date'))} {text(row.get('ClientVendor'))} {money(row.get('Receivable'))}", "")

    # Receivables and invoice log.
    src_recv, tgt_recv = group_receivables(src["Receivables"], "source"), group_receivables(tgt["Receivables"], "target")
    recv_source_only = sorted(set(src_recv) - set(tgt_recv))
    recv_target_only = sorted(set(tgt_recv) - set(src_recv))
    recv_mismatch = []
    for inv in sorted(set(src_recv) & set(tgt_recv)):
        s, t = src_recv[inv], tgt_recv[inv]
        diffs = []
        for field in ["total", "paid", "credits", "balance"]:
            if s[field] != t[field]:
                diffs.append(f"{field}: source ${s[field]:,.2f} vs CSPM ${t[field]:,.2f}")
        if norm_text(s["statuses"]) != norm_text(t["statuses"]):
            diffs.append(f"status: source {s['statuses']} vs CSPM {t['statuses']}")
        if diffs:
            recv_mismatch.append((inv, s, t, diffs))
    for inv in recv_source_only:
        s = src_recv[inv]
        add_record(records, "Receivables/A/R", "High", "Source receivable invoice not found in CSPM", ",".join(text(r.get("_row")) for r in s["rows"]), "", inv, f"total ${s['total']:,.2f}; balance ${s['balance']:,.2f}; status {s['statuses']}", "", "")
    for inv in recv_target_only:
        t = tgt_recv[inv]
        add_record(records, "Receivables/A/R", "Medium", "CSPM receivable invoice not found in source", "", ",".join(text(r.get("_row")) for r in t["rows"]), inv, "", f"total ${t['total']:,.2f}; balance ${t['balance']:,.2f}; status {t['statuses']}", "")
    for inv, s, t, diffs in recv_mismatch:
        add_record(records, "Receivables/A/R", "High", "Receivable invoice values/status differ", ",".join(text(r.get("_row")) for r in s["rows"]), ",".join(text(r.get("_row")) for r in t["rows"]), inv, f"total ${s['total']:,.2f}; paid ${s['paid']:,.2f}; balance ${s['balance']:,.2f}; status {s['statuses']}", f"total ${t['total']:,.2f}; paid ${t['paid']:,.2f}; balance ${t['balance']:,.2f}; status {t['statuses']}", "; ".join(diffs))

    src_invlog, tgt_invlog = group_invoice_log(src["Invoice Log"], "source"), group_invoice_log(tgt["InvoiceLog"], "target")
    invlog_source_only = sorted(set(src_invlog) - set(tgt_invlog))
    invlog_target_only = sorted(set(tgt_invlog) - set(src_invlog))
    invlog_mismatch = []
    for inv in sorted(set(src_invlog) & set(tgt_invlog)):
        s, t = src_invlog[inv], tgt_invlog[inv]
        diffs = []
        for field in ["fees", "disb", "tax", "aggregate"]:
            if s[field] != t[field]:
                diffs.append(f"{field}: source ${s[field]:,.2f} vs CSPM ${t[field]:,.2f}")
        if norm_text(s["client"]) != norm_text(t["client"]):
            diffs.append(f"client: source {s['client']} vs CSPM {t['client']}")
        if diffs:
            invlog_mismatch.append((inv, s, t, diffs))
    for inv in invlog_source_only:
        s = src_invlog[inv]
        add_record(records, "Invoice log", "High", "Source invoice-log invoice not found in CSPM", ",".join(text(r.get("_row")) for r in s["rows"]), "", inv, f"fees ${s['fees']:,.2f}; disb ${s['disb']:,.2f}; tax ${s['tax']:,.2f}; total ${s['aggregate']:,.2f}", "", "")
    for inv in invlog_target_only:
        t = tgt_invlog[inv]
        add_record(records, "Invoice log", "Medium", "CSPM invoice-log invoice not found in source", "", ",".join(text(r.get("_row")) for r in t["rows"]), inv, "", f"fees ${t['fees']:,.2f}; disb ${t['disb']:,.2f}; tax ${t['tax']:,.2f}; total ${t['aggregate']:,.2f}", "")
    for inv, s, t, diffs in invlog_mismatch:
        add_record(records, "Invoice log", "High", "Invoice-log values differ", ",".join(text(r.get("_row")) for r in s["rows"]), ",".join(text(r.get("_row")) for r in t["rows"]), inv, f"fees ${s['fees']:,.2f}; disb ${s['disb']:,.2f}; tax ${s['tax']:,.2f}; total ${s['aggregate']:,.2f}", f"fees ${t['fees']:,.2f}; disb ${t['disb']:,.2f}; tax ${t['tax']:,.2f}; total ${t['aggregate']:,.2f}", "; ".join(diffs))

    # Clients and matters.
    source_client_names = {norm_name(r.get("Client Name")): r for r in src["Clients"] if clean(r.get("Client Name"))}
    target_client_names = {norm_name(r.get("ClientName")): r for r in tgt["Clients"] if clean(r.get("ClientName"))}
    target_parent_names = {norm_name(r.get("ParentName")): r for r in tgt["Parents"] if clean(r.get("ParentName"))}
    client_missing = [r for key, r in source_client_names.items() if key not in target_client_names and key not in target_parent_names]
    client_extra = [r for key, r in target_client_names.items() if key not in source_client_names]
    source_client_ids = {clean(r.get("Client_ID")): r for r in src["Clients"] if clean(r.get("Client_ID"))}
    target_client_ids = {clean(r.get("ClientID")): r for r in tgt["Clients"] if clean(r.get("ClientID"))}
    target_parent_ids = {clean(r.get("ParentID")): r for r in tgt["Parents"] if clean(r.get("ParentID"))}
    client_id_missing = [cid for cid in source_client_ids if cid not in target_client_ids and cid not in target_parent_ids]
    for row in client_missing:
        add_record(records, "Clients", "Medium", "Source client name not represented as CSPM client or parent", row.get("_row"), "", row.get("Client Name"), row.get("Client Name"), "", "May be renamed/normalized in CSPM.")
    for cid in client_id_missing:
        row = source_client_ids[cid]
        add_record(records, "Clients", "Medium", "Source Client_ID not represented as CSPM ClientID or ParentID", row.get("_row"), "", cid, f"{cid} / {text(row.get('Client Name'))}", "", "")
    for row in client_extra:
        add_record(records, "Clients", "Low", "CSPM client name not found in source Clients sheet", "", row.get("_row"), row.get("ClientName"), "", row.get("ClientName"), "Often expected because CSPM splits parent/work clients and includes newer clients.")

    source_matters = {clean(r.get("Matter_ID")): r for r in src["Matters"] if clean(r.get("Matter_ID"))}
    target_matters = {clean(r.get("MatterNumber")): r for r in tgt["Matters"] if clean(r.get("MatterNumber"))}
    matter_missing = [source_matters[k] for k in sorted(set(source_matters) - set(target_matters))]
    matter_extra = [target_matters[k] for k in sorted(set(target_matters) - set(source_matters))]
    for row in matter_missing:
        add_record(records, "Matters", "High", "Source matter not found in CSPM Matters", row.get("_row"), "", row.get("Matter_ID"), row.get("Description (Internal)"), "", "Docket rows tied to this matter may lose matter linkage.")
    for row in matter_extra:
        add_record(records, "Matters", "Low", "CSPM matter not found in source Matters", "", row.get("_row"), row.get("MatterNumber"), "", row.get("MatterName"), "Likely newer CSPM matter or generated during import.")

    if not tgt["Transactions"]:
        add_record(
            records,
            "Transactions canonical table",
            "High",
            "CSPM Transactions sheet is empty while legacy finance sheets contain ledger/A/R/invoice data",
            "",
            "",
            "Transactions",
            "Source has Ledger, Receivables, Invoice Log, and Disbursements",
            "Transactions rows: 0",
            "If finance screens expect Transactions as canonical, financial history is missing there.",
        )

    def summarize_recv(groups):
        return {
            "Invoices": len(groups),
            "Rows": sum(len(g["rows"]) for g in groups.values()),
            "TotalInvoiced": sum(g["total"] for g in groups.values()),
            "AmountPaid": sum(g["paid"] for g in groups.values()),
            "CreditsAdj": sum(g["credits"] for g in groups.values()),
            "BalanceDue": sum(g["balance"] for g in groups.values()),
        }

    def summarize_invlog(groups):
        return {
            "Invoices": len(groups),
            "Rows": sum(len(g["rows"]) for g in groups.values()),
            "Fees": sum(g["fees"] for g in groups.values()),
            "Disbursements": sum(g["disb"] for g in groups.values()),
            "Tax": sum(g["tax"] for g in groups.values()),
            "Aggregate": sum(g["aggregate"] for g in groups.values()),
        }

    source_dockets = summarize_dockets(src["Dockets"], "source")
    target_dockets = summarize_dockets(tgt["Dockets"], "target")
    target_time = summarize_dockets(tgt["TimeEntries"], "time")
    summary = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source": str(SOURCE),
        "target": str(TARGET),
        "sqlite_db": str(SQLITE),
        "sqlite_db_size_bytes": SQLITE.stat().st_size if SQLITE.exists() else None,
        "source_modified": datetime.fromtimestamp(SOURCE.stat().st_mtime).isoformat(timespec="seconds"),
        "target_modified": datetime.fromtimestamp(TARGET.stat().st_mtime).isoformat(timespec="seconds"),
        "row_counts": {"source": {k: len(v) for k, v in src.items()}, "cspm": {k: len(v) for k, v in tgt.items()}},
        "dockets": {
            "source": source_dockets,
            "cspm_legacy_dockets": target_dockets,
            "cspm_timeentries": target_time,
            "row_delta": target_dockets["Rows"] - source_dockets["Rows"],
            "hour_delta": target_dockets["Hours"] - source_dockets["Hours"],
            "amount_to_cs_delta": target_dockets["AmountToCS"] - source_dockets["AmountToCS"],
            "total_incl_hst_delta": target_dockets["TotalInclHST"] - source_dockets["TotalInclHST"],
            "source_missing_by_core": len(missing_dockets),
            "cspm_extra_by_core": len(extra_dockets),
            "invoice_status_changes_on_matched_rows": len(invoice_changes),
            "client_label_changes_on_matched_rows": len(client_label_changes),
            "timeentries_source_missing_by_core": sum((time_source_counter - time_target_counter).values()),
            "timeentries_cspm_extra_by_core": sum((time_target_counter - time_source_counter).values()),
        },
        "disbursements": {
            "source": summarize_disb(src["Disbursements"], "source"),
            "cspm": summarize_disb(tgt["Disbursements"], "target"),
            "source_missing": len(disb_missing),
            "cspm_extra": len(disb_extra),
        },
        "ledger": {
            "source": summarize_ledger(src["Ledger"], "source"),
            "cspm": summarize_ledger(tgt["Ledger"], "target"),
            "mismatches": len(ledger_mismatch),
            "source_only": len(ledger_source_only),
            "cspm_only": len(ledger_target_only),
            "malformed_source_rows_excluded": len(malformed_source_ledger),
        },
        "receivables": {
            "source": summarize_recv(src_recv),
            "cspm": summarize_recv(tgt_recv),
            "source_only_invoices": len(recv_source_only),
            "cspm_only_invoices": len(recv_target_only),
            "mismatched_invoices": len(recv_mismatch),
            "source_rollup_rows_excluded": len(source_receivable_rollups),
            "cspm_rollup_rows_excluded": len(target_receivable_rollups),
        },
        "invoice_log": {
            "source": summarize_invlog(src_invlog),
            "cspm": summarize_invlog(tgt_invlog),
            "source_only_invoices": len(invlog_source_only),
            "cspm_only_invoices": len(invlog_target_only),
            "mismatched_invoices": len(invlog_mismatch),
        },
        "clients": {
            "source_client_names": len(source_client_names),
            "cspm_client_names": len(target_client_names),
            "cspm_parent_names": len(target_parent_names),
            "source_names_not_represented": len(client_missing),
            "source_ids_not_represented": len(client_id_missing),
            "cspm_client_names_not_in_source": len(client_extra),
        },
        "matters": {
            "source_matters": len(source_matters),
            "cspm_matters": len(target_matters),
            "source_missing_in_cspm": len(matter_missing),
            "cspm_extra": len(matter_extra),
        },
        "discrepancy_records": len(records),
    }
    return summary, records


def serializable(value):
    if isinstance(value, Decimal):
        return f"{value:,.2f}"
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    return text(value)


def md_escape(value) -> str:
    return text(value).replace("|", "\\|").replace("\n", " ")


def md_table(headers, rows) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(md_escape(v) for v in row) + " |")
    return "\n".join(lines)


def write_outputs(summary: dict, records: list[dict]) -> None:
    with CSV_PATH.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=["Category", "Severity", "Issue", "SourceRow", "CSPMRow", "Identifier", "SourceValue", "CSPMValue", "Note"])
        writer.writeheader()
        writer.writerows(records)
    JSON_PATH.write_text(json.dumps(summary, indent=2, default=serializable), encoding="utf-8")

    dockets = summary["dockets"]
    lines = [
        "# Dockets.xlsm to CSPM Reconciliation Report",
        "",
        f"Generated: {summary['generated_at']}",
        f"Source of truth: `{summary['source']}`",
        f"Compared CSPM store: `{summary['target']}`",
        f"SQLite database checked: `{summary['sqlite_db']}` ({summary['sqlite_db_size_bytes']} bytes)",
        "",
        "## Executive Summary",
        "",
        "- The repo-root `CSPM_Database.db` is 0 bytes. A literal SQLite comparison would find no app database rows. This report uses the populated `data/CSPM.xlsm` workbook as the practical CSPM data store.",
        f"- Docketed time/fees do not reconcile exactly: {dockets['source_missing_by_core']} source rows are not found in CSPM by core time/fee fields; {dockets['cspm_extra_by_core']} CSPM rows are not present in the current source workbook.",
        f"- CSPM has {dockets['row_delta']} more docket rows, {dockets['hour_delta']} more hours, {money(dockets['amount_to_cs_delta'])} more Amount to CS, and {money(dockets['total_incl_hst_delta'])} more total incl. HST than the source workbook.",
        f"- {dockets['invoice_status_changes_on_matched_rows']} matched source docket rows have different invoice/WIP markers in CSPM.",
        f"- Disbursements reconcile by normalized core fields: {summary['disbursements']['source_missing']} source-only and {summary['disbursements']['cspm_extra']} CSPM-only rows.",
        f"- Receivables/A/R diverge: {summary['receivables']['mismatched_invoices']} shared invoice numbers differ; {summary['receivables']['cspm_only_invoices']} CSPM receivable invoices are not in source.",
        f"- Invoice Log diverges: {summary['invoice_log']['mismatched_invoices']} shared invoice numbers differ; {summary['invoice_log']['cspm_only_invoices']} CSPM invoice-log invoices are not in source.",
        f"- Ledger diverges: {summary['ledger']['mismatches']} shared key mismatch, {summary['ledger']['source_only']} source-only rows, {summary['ledger']['cspm_only']} CSPM-only rows, and {summary['ledger']['malformed_source_rows_excluded']} malformed source rows excluded from keyed comparison.",
        "- The CSPM `Transactions` sheet is empty. This is a high-risk architecture discrepancy if any finance screen/report expects Transactions to be canonical.",
        "",
        "## Row Counts",
        md_table(["Sheet", "Dockets.xlsm", "CSPM.xlsm"], [[k, summary["row_counts"]["source"].get(k, ""), summary["row_counts"]["cspm"].get(k, "")] for k in sorted(set(summary["row_counts"]["source"]) | set(summary["row_counts"]["cspm"]))]),
        "",
        "## Financial Totals",
        md_table(
            ["Domain", "Source", "CSPM", "Delta / note"],
            [
                ["Docket rows", summary["dockets"]["source"]["Rows"], summary["dockets"]["cspm_legacy_dockets"]["Rows"], summary["dockets"]["row_delta"]],
                ["Docket hours", summary["dockets"]["source"]["Hours"], summary["dockets"]["cspm_legacy_dockets"]["Hours"], summary["dockets"]["hour_delta"]],
                ["Docket Amount to CS", money(summary["dockets"]["source"]["AmountToCS"]), money(summary["dockets"]["cspm_legacy_dockets"]["AmountToCS"]), money(summary["dockets"]["amount_to_cs_delta"])],
                ["Docket total incl. HST", money(summary["dockets"]["source"]["TotalInclHST"]), money(summary["dockets"]["cspm_legacy_dockets"]["TotalInclHST"]), money(summary["dockets"]["total_incl_hst_delta"])],
                ["Disbursement amount", money(summary["disbursements"]["source"]["Amount"]), money(summary["disbursements"]["cspm"]["Amount"]), ""],
                ["Invoice log aggregate", money(summary["invoice_log"]["source"]["Aggregate"]), money(summary["invoice_log"]["cspm"]["Aggregate"]), ""],
                ["Receivables balance due", money(summary["receivables"]["source"]["BalanceDue"]), money(summary["receivables"]["cspm"]["BalanceDue"]), ""],
                ["Ledger receivable sum", money(summary["ledger"]["source"]["Receivable"]), money(summary["ledger"]["cspm"]["Receivable"]), ""],
            ],
        ),
        "",
        "## Assumptions And Uncertainties",
        "",
        "- `Dockets.xlsm` is treated as source of truth.",
        "- `data/CSPM.xlsm` is treated as the populated CSPM app data store because the SQLite DB is empty.",
        "- Time-entry matching ignores client-name normalization and invoice/status labels first, then reports those separately.",
        "- Actual invoice numbers are values like `25-0053` or `26-0018-A`. `BILLED`, hold/forgot/free/write-off markers, and blanks are status/WIP/no-bill markers.",
        "- CSPM-only rows may be legitimate later work. If `Dockets.xlsm` must remain sole truth, those rows need to be backported or consciously rejected.",
        "",
        "## Discrepancy Detail",
        md_table(["Category", "Severity", "Issue", "Source row", "CSPM row", "Identifier", "Source", "CSPM", "Note"], [[r["Category"], r["Severity"], r["Issue"], r["SourceRow"], r["CSPMRow"], r["Identifier"], r["SourceValue"], r["CSPMValue"], r["Note"]] for r in records]),
        "",
        "## Artifact Files",
        "",
        f"- DOCX: `{DOCX_PATH}`",
        f"- CSV: `{CSV_PATH}`",
        f"- JSON: `{JSON_PATH}`",
    ]
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")

    if Document is None:
        return
    doc = Document()
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9)
    doc.add_heading("Dockets.xlsm to CSPM Reconciliation Report", 0)
    for item in [
        f"Generated: {summary['generated_at']}",
        f"Source of truth: {summary['source']}",
        f"Compared CSPM store: {summary['target']}",
        f"SQLite database checked: {summary['sqlite_db']} ({summary['sqlite_db_size_bytes']} bytes)",
    ]:
        doc.add_paragraph(item)
    doc.add_heading("Executive Summary", level=1)
    for item in lines[8:17]:
        if item.startswith("- "):
            doc.add_paragraph(item[2:], style="List Bullet")
    doc.add_heading("Financial Totals", level=1)
    add_docx_table(
        doc,
        ["Domain", "Source", "CSPM", "Delta / note"],
        [
            ["Docket rows", summary["dockets"]["source"]["Rows"], summary["dockets"]["cspm_legacy_dockets"]["Rows"], summary["dockets"]["row_delta"]],
            ["Docket hours", summary["dockets"]["source"]["Hours"], summary["dockets"]["cspm_legacy_dockets"]["Hours"], summary["dockets"]["hour_delta"]],
            ["Docket Amount to CS", money(summary["dockets"]["source"]["AmountToCS"]), money(summary["dockets"]["cspm_legacy_dockets"]["AmountToCS"]), money(summary["dockets"]["amount_to_cs_delta"])],
            ["Docket total incl. HST", money(summary["dockets"]["source"]["TotalInclHST"]), money(summary["dockets"]["cspm_legacy_dockets"]["TotalInclHST"]), money(summary["dockets"]["total_incl_hst_delta"])],
            ["Disbursement amount", money(summary["disbursements"]["source"]["Amount"]), money(summary["disbursements"]["cspm"]["Amount"]), ""],
            ["Invoice log aggregate", money(summary["invoice_log"]["source"]["Aggregate"]), money(summary["invoice_log"]["cspm"]["Aggregate"]), ""],
            ["Receivables balance due", money(summary["receivables"]["source"]["BalanceDue"]), money(summary["receivables"]["cspm"]["BalanceDue"]), ""],
            ["Ledger receivable sum", money(summary["ledger"]["source"]["Receivable"]), money(summary["ledger"]["cspm"]["Receivable"]), ""],
        ],
    )
    doc.add_heading("Key Counts", level=1)
    add_docx_table(
        doc,
        ["Area", "Count"],
        [
            ["Source docket rows missing by core fields", summary["dockets"]["source_missing_by_core"]],
            ["CSPM docket rows not in source", summary["dockets"]["cspm_extra_by_core"]],
            ["Matched docket rows with invoice/status changes", summary["dockets"]["invoice_status_changes_on_matched_rows"]],
            ["Disbursement source-only rows", summary["disbursements"]["source_missing"]],
            ["Disbursement CSPM-only rows", summary["disbursements"]["cspm_extra"]],
            ["Ledger mismatches", summary["ledger"]["mismatches"]],
            ["Ledger source-only rows", summary["ledger"]["source_only"]],
            ["Ledger CSPM-only rows", summary["ledger"]["cspm_only"]],
            ["Malformed source ledger rows excluded", summary["ledger"]["malformed_source_rows_excluded"]],
            ["Receivable mismatched invoices", summary["receivables"]["mismatched_invoices"]],
            ["Receivable CSPM-only invoices", summary["receivables"]["cspm_only_invoices"]],
            ["Invoice-log mismatched invoices", summary["invoice_log"]["mismatched_invoices"]],
            ["Invoice-log CSPM-only invoices", summary["invoice_log"]["cspm_only_invoices"]],
            ["Source client names not represented", summary["clients"]["source_names_not_represented"]],
            ["Source matters missing in CSPM", summary["matters"]["source_missing_in_cspm"]],
        ],
    )
    doc.add_heading("Discrepancy Detail", level=1)
    add_docx_table(
        doc,
        ["Category", "Severity", "Issue", "Source row", "CSPM row", "Identifier", "Source", "CSPM", "Note"],
        [[r["Category"], r["Severity"], r["Issue"], r["SourceRow"], r["CSPMRow"], r["Identifier"], r["SourceValue"], r["CSPMValue"], r["Note"]] for r in records],
    )
    doc.add_heading("Artifact Files", level=1)
    for path in [MD_PATH, CSV_PATH, JSON_PATH]:
        doc.add_paragraph(str(path), style="List Bullet")
    doc.save(DOCX_PATH)


def add_docx_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = text(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)


def main() -> int:
    summary, records = compare()
    write_outputs(summary, records)
    print(f"DOCX={DOCX_PATH if DOCX_PATH.exists() else ''}")
    print(f"MD={MD_PATH}")
    print(f"CSV={CSV_PATH}")
    print(f"JSON={JSON_PATH}")
    print(
        json.dumps(
            {
                "docket_source_missing": summary["dockets"]["source_missing_by_core"],
                "docket_cspm_extra": summary["dockets"]["cspm_extra_by_core"],
                "docket_invoice_changes": summary["dockets"]["invoice_status_changes_on_matched_rows"],
                "disbursement_source_missing": summary["disbursements"]["source_missing"],
                "disbursement_cspm_extra": summary["disbursements"]["cspm_extra"],
                "ledger_mismatches": summary["ledger"]["mismatches"],
                "ledger_source_only": summary["ledger"]["source_only"],
                "ledger_cspm_only": summary["ledger"]["cspm_only"],
                "ledger_malformed_source_rows": summary["ledger"]["malformed_source_rows_excluded"],
                "receivable_mismatches": summary["receivables"]["mismatched_invoices"],
                "receivable_cspm_only": summary["receivables"]["cspm_only_invoices"],
                "invoice_log_mismatches": summary["invoice_log"]["mismatched_invoices"],
                "invoice_log_cspm_only": summary["invoice_log"]["cspm_only_invoices"],
                "records": summary["discrepancy_records"],
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
